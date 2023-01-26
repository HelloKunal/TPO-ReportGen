from get_data_from_sheets import GetValues
from docx import Document
from docx.shared import Inches
import numpy
import matplotlib.pyplot as matpyplot
from textwrap import wrap
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from datetime import date
import json

def generate_docx(branch):
    document = Document('static/default_theme.docx')

    document.add_picture('static/front_page.png', width=Inches(6), height=Inches(8))

    get_values_object = GetValues(branch)
    data = get_values_object.compl_proc()
    document.add_page_break()
    # generate_docs_object.createNew()
    # fte_data = fte_data[:10]
    # print(fte_data)
    document.add_heading(branch+' DEPARTMENT (BATCH 2022-2023)', level=0)
    FIELDS = [
        'No. of companies visited',
        'Total Strength',
        'No. of offers',
        'No. of students placed',
        'No. of students with 6-month Internship',
        'Percentage Placement',
        'Average Package',
        'Median Package',
        'Highest Package',
        'Lowest Package',
    ]      
    document.add_heading('1.    Placement Stats 2022-23:', level=1)
    table = document.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    special_values = data[3]
    for i in range(0, 10):
        row_cells = table.add_row().cells
        row_cells[0].text = FIELDS[i]
        row_cells[1].text = str(special_values[i])
    # para = document.add_paragraph()
    # run = para.add_run()
    # run.add_break(WD_BREAK.LINE)


    graph_values = special_values[:5]
    graph_fields = FIELDS[:5]
    graph_fields = [ '\n'.join(wrap(l, 20)) for l in graph_fields ]
    fig = matpyplot.figure(figsize = (10, 5))
    matpyplot.bar(graph_fields, graph_values, color ='#4472C4',
            width = 0.4)
    for i in range(len(graph_fields)):
        matpyplot.text(i, graph_values[i], graph_values[i], ha = 'center')
    matpyplot.title("Placement Statistics")
    
    # matpyplot.xlabel("Courses offered")
    # matpyplot.ylabel("No. of students enrolled")
    # matpyplot.title("Students enrolled in different courses")
    # matpyplot.show()
    matpyplot.savefig('special_values_bar.png')
    document.add_picture('special_values_bar.png', width=Inches(5), height=Inches(2.5))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # para = document.add_paragraph()
    # run = para.add_run()
    # run.add_break(WD_BREAK.LINE)

    pie_fig = matpyplot.figure(figsize =(10, 7))
    pie_values = [special_values[3], special_values[1] - special_values[3]]
    pie_fields = ['Placed Students', 'Unplaced Students']
    pie_explode = (0, 0.2)
    pie_colors = ["#4472C4", "#D9E2F3"]
    matpyplot.pie(pie_values, labels = pie_fields, autopct=lambda p: '{:.0f}%'.format(p), shadow=True, startangle=90, explode=pie_explode, colors=pie_colors)
    matpyplot.title("Placed and unplaced students")
    matpyplot.savefig('special_values_pie.png')
    document.add_picture('special_values_pie.png', width=Inches(5), height=Inches(2.5))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()


    document.add_heading('1.1   Campus Placement List (FTE):', level=1)
    fte_data = data[0]
    table = document.add_table(rows=1, cols=len(fte_data[0]))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'S. No'
    hdr_cells[1].text = 'Scholar Number'
    hdr_cells[2].text = 'Candidate Name'
    hdr_cells[3].text = "Company's Name"
    hdr_cells[4].text = 'CTC (LPA)'
    for i in range(len(fte_data)):
        row_cells = table.add_row().cells
        row_cells[0].text = str(fte_data[i][0])
        row_cells[1].text = str(fte_data[i][1])
        row_cells[2].text = str(fte_data[i][2])
        row_cells[3].text = str(fte_data[i][3])
        row_cells[4].text = str(fte_data[i][4])

    document.add_heading('1.2   Students With PPO:', level=1)
    ppo_data = data[1]
    table = document.add_table(rows=1, cols=len(ppo_data[0]))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'S. No'
    hdr_cells[1].text = 'Scholar Number'
    hdr_cells[2].text = 'Candidate Name'
    hdr_cells[3].text = "Company's Name"
    hdr_cells[4].text = 'CTC (LPA)'
    for i in range(len(ppo_data)):
        row_cells = table.add_row().cells
        row_cells[0].text = str(ppo_data[i][0])
        row_cells[1].text = str(ppo_data[i][1])
        row_cells[2].text = str(ppo_data[i][2])
        row_cells[3].text = str(ppo_data[i][3])
        row_cells[4].text = str(ppo_data[i][4])

    document.add_heading('1.2   Internships (6 Month):', level=1)
    sixm_data = data[2]
    table = document.add_table(rows=1, cols=len(sixm_data[0]))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'S. No'
    hdr_cells[1].text = 'Scholar Number'
    hdr_cells[2].text = 'Candidate Name'
    hdr_cells[3].text = "Company's Name"
    hdr_cells[4].text = 'Stipend (KPM)'
    for i in range(len(sixm_data)):
        row_cells = table.add_row().cells
        row_cells[0].text = str(sixm_data[i][0])
        row_cells[1].text = str(sixm_data[i][1])
        row_cells[2].text = str(sixm_data[i][2])
        row_cells[3].text = str(sixm_data[i][3])
        row_cells[4].text = str(sixm_data[i][4])
    document.add_page_break()

    document.add_heading('1.4   Companies visited:', level=1)
    company_names = data[4]
    table = document.add_table(rows=0, cols=4)
    table.style = 'Table Grid'
    for i in range(0, len(company_names) // 4):
        row_cells = table.add_row().cells
        row_cells[0].text = company_names[4*i]
        row_cells[1].text = company_names[4*i+1]
        row_cells[2].text = company_names[4*i+2]
        row_cells[3].text = company_names[4*i+3]

    document.add_heading('1.5.	Companies visited(with selections):', level=1)
    companies_with_selections = data[6]
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'S. No'
    hdr_cells[1].text = 'Company Name'
    hdr_cells[2].text = 'Selections'
    hdr_cells[3].text = 'CTC (LPA)'
    for i in range(len(companies_with_selections)):
        row_cells = table.add_row().cells
        row_cells[0].text = str(companies_with_selections[i][0])
        row_cells[1].text = str(companies_with_selections[i][1])
        row_cells[2].text = str(companies_with_selections[i][2])
        row_cells[3].text = str(companies_with_selections[i][3])

    document.add_heading('1.6.	Companies visited(with no selections):', level=1)
    branch_companies_without_selections = data[5]
    table = document.add_table(rows=0, cols=4)
    table.style = 'Table Grid'
    for i in range(0, len(branch_companies_without_selections) // 4):
        row_cells = table.add_row().cells
        row_cells[0].text = branch_companies_without_selections[4*i]
        row_cells[1].text = branch_companies_without_selections[4*i+1]
        row_cells[2].text = branch_companies_without_selections[4*i+2]
        row_cells[3].text = branch_companies_without_selections[4*i+3]

    with open('static/prev_year_data.json') as f:
        file_string = f.read()
    PREV_YEAR_DATA = json.loads(file_string)
    PREV_YEAR_DATA = PREV_YEAR_DATA[branch]

    if len(PREV_YEAR_DATA):
        document.add_page_break()
        PREV_YEAR_DATA.append(["2022-23", special_values[0], special_values[2], special_values[6]])
        document.add_heading('2.1   Tabular Comparison:', level=1)
        table = document.add_table(rows=4, cols=1)
        table.style = 'Table Grid'
        hdr_cells = table.columns[0].cells
        hdr_cells[0].text = 'Session'
        hdr_cells[1].text = 'No of Companies visited'
        hdr_cells[2].text = 'No. of oﬀers'
        hdr_cells[3].text = 'Average Package'
        for i in range(len(PREV_YEAR_DATA)):
            col_cells = table.add_column(Inches(3)).cells
            col_cells[0].text = str(PREV_YEAR_DATA[i][0])
            col_cells[1].text = str(PREV_YEAR_DATA[i][1])
            col_cells[2].text = str(PREV_YEAR_DATA[i][2])
            col_cells[3].text = str(PREV_YEAR_DATA[i][3])


        document.add_heading('2.2   Graphs for Comparison with previous years:', level=1)
        document.add_heading('a)	 No of companies visited:', level=2)
        graph_values = [int(i[1]) for i in PREV_YEAR_DATA]
        graph_fields = [i[0] for i in PREV_YEAR_DATA]
        fig = matpyplot.figure(figsize = (10, 5))
        matpyplot.bar(graph_fields, graph_values, color ='#4472C4',
                width = 0.4)
        matpyplot.title("No of companies visited")
        for i in range(len(graph_fields)):
            matpyplot.text(i, graph_values[i], graph_values[i], ha = 'center')
        matpyplot.savefig('prev_no_companies.png')
        document.add_picture('prev_no_companies.png', width=Inches(6), height=Inches(3))

        document.add_heading('b)	 Average package distribution:', level=2)
        graph_values = [int(i[3]) for i in PREV_YEAR_DATA]
        graph_fields = [i[0] for i in PREV_YEAR_DATA]
        fig = matpyplot.figure(figsize = (10, 5))
        matpyplot.bar(graph_fields, graph_values, color ='#4472C4',
                width = 0.4)
        matpyplot.title("Average Package")
        for i in range(len(graph_fields)):
            matpyplot.text(i, graph_values[i], graph_values[i], ha = 'center')
        matpyplot.savefig('prev_avg_package.png')
        document.add_picture('prev_avg_package.png', width=Inches(6), height=Inches(3))
                
        document.add_heading('c)	No. of Offers:', level=2)
        graph_values = [int(i[2]) for i in PREV_YEAR_DATA]
        graph_fields = [i[0] for i in PREV_YEAR_DATA]
        fig = matpyplot.figure(figsize = (10, 5))
        matpyplot.bar(graph_fields, graph_values, color ='#4472C4',
                width = 0.4)
        matpyplot.title("No. of offers")
        for i in range(len(graph_fields)):
            matpyplot.text(i, graph_values[i], graph_values[i], ha = 'center')
        matpyplot.savefig('prev_no_offers.png')
        document.add_picture('prev_no_offers.png', width=Inches(6), height=Inches(3))

    today = date.today()
    curr_timestamp = today.strftime("%B_%d_%Y")
    document.save('static/TPO_REPORT_'+branch+'_'+curr_timestamp+'.docx')